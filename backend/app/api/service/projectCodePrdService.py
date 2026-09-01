# encoding: UTF-8
import os
import re
import subprocess
import threading
from io import BytesIO
from pathlib import Path


from logger import logger
from common.sqlSession import SqlSession
from ..dao.projectCodePrdDao import ProjectCodePrdDao
from ..model.productModel import Product
from ..model.projectModel import Project
from ..model.projectCodePrdModel import ProjectCodePrdConfig, ProjectCodePrdRecord
from .aiService import AIService


class ProjectCodePrdService(object):
    @staticmethod
    def create(session, model_cls, add_info):
        return ProjectCodePrdDao.create(session, model_cls, add_info)

    @staticmethod
    def update_by_id(session, model_cls, obj_id, update_info, soft_delete=True):
        return ProjectCodePrdDao.update_by_id(session, model_cls, obj_id, update_info, soft_delete)

    @staticmethod
    def get_by_id(session, model_cls, obj_id, soft_delete=True):
        return ProjectCodePrdDao.get_by_id(session, model_cls, obj_id, soft_delete)

    @staticmethod
    def get_config_by_project(session, project_id):
        return ProjectCodePrdDao.first_by_filters(
            session,
            ProjectCodePrdConfig,
            [ProjectCodePrdConfig.project_id == int(project_id)]
        )

    @staticmethod
    def start_generate_prd(record_id, prompt_append=''):
        def run():
            session = SqlSession()
            try:
                ProjectCodePrdService.generate_prd(session, record_id, prompt_append)
            except Exception as e:
                logger.error(f'后台生成代码PRD失败: {e}')
                try:
                    ProjectCodePrdService.update_by_id(session, ProjectCodePrdRecord, record_id, {
                        'status': 3,
                        'error_message': f'后台生成代码PRD失败: {e}'
                    })
                except Exception as update_err:
                    logger.error(f'更新代码PRD失败状态失败: {update_err}')
            finally:
                session.close()
        threading.Thread(target=run, daemon=True).start()

    @staticmethod
    def list_records(session, project_id, page_no=1, page_size=10):
        return ProjectCodePrdDao.list_by_filters(
            session,
            ProjectCodePrdRecord,
            [ProjectCodePrdRecord.project_id == int(project_id)],
            int(page_no),
            int(page_size),
            ProjectCodePrdRecord.created_time
        )

    @staticmethod
    def list_remote_branches(repo_url):
        if not ProjectCodePrdService._is_safe_git_url(repo_url):
            return [], 'Git仓库地址格式不正确'
        try:
            proc = subprocess.run(
                ['git', 'ls-remote', '--heads', repo_url],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=30,
                shell=False
            )
            if proc.returncode != 0:
                return [], ProjectCodePrdService._format_process_output(proc, '获取分支失败', 500)

            branches = []
            for line in proc.stdout.splitlines():
                if 'refs/heads/' not in line:
                    continue
                branches.append(line.rsplit('refs/heads/', 1)[-1].strip())
            return sorted(set(branches)), ''
        except Exception as e:
            logger.error(f'获取Git分支失败: {e}')
            return [], f'获取分支失败: {e}'

    @staticmethod
    def generate_prd(session, record_id, prompt_append=''):
        record = ProjectCodePrdService.get_by_id(session, ProjectCodePrdRecord, record_id)
        if not record:
            return 0, '未查询到生成记录'
        ProjectCodePrdService.update_by_id(session, ProjectCodePrdRecord, record.id, {'status': 1, 'error_message': ''})
        try:
            summary, err_msg = ProjectCodePrdService.build_repo_summary(session, record.project_id, record.repo_url, record.branch)
            if err_msg:
                ProjectCodePrdService.update_by_id(session, ProjectCodePrdRecord, record.id, {'status': 3, 'error_message': err_msg})
                return 0, err_msg
            markdown, err_msg = ProjectCodePrdService.generate_prd_markdown(summary, record.repo_url, record.branch, prompt_append)
            if err_msg:
                ProjectCodePrdService.update_by_id(session, ProjectCodePrdRecord, record.id, {
                    'status': 3,
                    'summary': summary,
                    'error_message': err_msg
                })
                return 0, err_msg
            title = ProjectCodePrdService.extract_title(markdown) or f'代码转PRD-{record.branch}'
            return ProjectCodePrdService.update_by_id(session, ProjectCodePrdRecord, record.id, {
                'status': 2,
                'summary': summary,
                'prd_markdown': markdown,
                'title': title,
                'error_message': ''
            })
        except Exception as e:
            logger.error(f'生成代码PRD失败: {e}')
            ProjectCodePrdService.update_by_id(session, ProjectCodePrdRecord, record.id, {'status': 3, 'error_message': str(e)})
            return 0, f'生成代码PRD失败: {e}'

    @staticmethod
    def build_repo_summary(session, project_id, repo_url, branch):
        if not ProjectCodePrdService._is_safe_git_url(repo_url):
            return '', 'Git仓库地址格式不正确'
        repo_dir = ProjectCodePrdService._get_repo_work_dir(session, project_id, branch)
        err_msg = ProjectCodePrdService._sync_repo(repo_url, branch, repo_dir)
        if err_msg:
            return '', err_msg
        return ProjectCodePrdService._scan_repo(str(repo_dir), repo_url, branch), ''

    @staticmethod
    def _get_repo_work_dir(session, project_id, branch):
        project = session.query(Project).filter(Project.id == int(project_id)).first() if project_id else None
        product = None
        if project and project.product_id:
            product = session.query(Product).filter(Product.id == int(project.product_id)).first()
        product_name = ProjectCodePrdService._safe_path_name(product.name if product else '未归属产品')
        project_name = ProjectCodePrdService._safe_path_name(project.name if project else f'project-{project_id}')
        branch_name = ProjectCodePrdService._safe_path_name(branch or 'default')
        repo_root = Path(__file__).resolve().parents[3] / 'downloads-git'
        return repo_root / product_name / f'{project_name}-{branch_name}'

    @staticmethod
    def _sync_repo(repo_url, branch, repo_dir):
        repo_dir = Path(repo_dir)
        repo_dir.parent.mkdir(parents=True, exist_ok=True)
        if not repo_dir.exists():
            return ProjectCodePrdService._run_git(['git', 'clone', '--depth', '1', '--branch', branch, repo_url, str(repo_dir)], None, '拉取仓库失败', 600)
        if not (repo_dir / '.git').exists():
            return f'本地目录已存在但不是Git仓库: {repo_dir}'
        err_msg = ProjectCodePrdService._run_git(['git', 'remote', 'set-url', 'origin', repo_url], str(repo_dir), '更新仓库地址失败', 60)
        if err_msg:
            return err_msg
        err_msg = ProjectCodePrdService._run_git(['git', 'fetch', 'origin', branch, '--depth', '1', '--prune'], str(repo_dir), '更新远端代码失败', 600)
        if err_msg:
            return err_msg
        err_msg = ProjectCodePrdService._run_git(['git', 'checkout', '-B', branch, f'origin/{branch}'], str(repo_dir), '切换分支失败', 120)
        if err_msg:
            return err_msg
        return ProjectCodePrdService._run_git(['git', 'reset', '--hard', f'origin/{branch}'], str(repo_dir), '更新工作区失败', 120)

    @staticmethod
    def _run_git(cmd, cwd, error_prefix, timeout):
        try:
            proc = subprocess.run(
                cmd,
                cwd=cwd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=timeout,
                shell=False
            )
            if proc.returncode != 0:
                return f"{error_prefix}: {ProjectCodePrdService._format_process_output(proc, 'Git命令执行失败', 1000)}"

            return ''
        except Exception as e:
            logger.error(f'{error_prefix}: {e}')
            return f'{error_prefix}: {e}'

    @staticmethod
    def _format_process_output(proc, default_message, limit):
        output = (getattr(proc, 'stderr', '') or getattr(proc, 'stdout', '') or default_message).strip()
        return output[:limit]

    @staticmethod
    def generate_prd_markdown(summary, repo_url, branch, prompt_append=''):

        analysis_summary, err_msg = ProjectCodePrdService._build_prd_analysis_summary(summary, repo_url, branch, prompt_append)
        if err_msg:
            return '', err_msg
        prompt = ProjectCodePrdService._build_prd_prompt(analysis_summary, repo_url, branch, prompt_append)
        result, err_msg = ProjectCodePrdService._request_markdown(prompt, max_tokens=7000)
        if err_msg:
            return '', err_msg
        markdown = ProjectCodePrdService._strip_markdown_fence(result)
        if not markdown.strip():
            return '', '大模型未返回PRD内容'
        return markdown, ''

    @staticmethod
    def _build_prd_analysis_summary(summary, repo_url, branch, prompt_append=''):
        chunks = ProjectCodePrdService._split_summary_chunks(summary)
        if len(chunks) <= 1:
            return summary, ''
        results = []
        logger.info(f'代码PRD启用串行分块agent分析: chunks={len(chunks)}')
        try:
            for index, chunk in enumerate(chunks):
                chunk_result, err_msg = ProjectCodePrdService._analyze_prd_chunk(
                    chunk,
                    index + 1,
                    len(chunks),
                    repo_url,
                    branch,
                    prompt_append
                )
                if err_msg:
                    return '', err_msg
                results.append(chunk_result)
        except Exception as e:
            logger.error(f'代码PRD分块agent分析失败: {e}')
            return '', f'代码PRD分块agent分析失败: {e}'
        merged = '\n\n'.join([item for item in results if item])
        if not merged.strip():
            return '', '代码PRD分块agent未返回分析内容'
        return f'''
以下是代码库分块agent分析结果，请基于这些结果生成最终PRD，不要再逐文件展开原始代码。

{merged}
'''.strip(), ''

    @staticmethod
    def _analyze_prd_chunk(chunk, chunk_index, total_chunks, repo_url, branch, prompt_append=''):
        prompt = ProjectCodePrdService._build_prd_chunk_prompt(chunk, chunk_index, total_chunks, repo_url, branch, prompt_append)
        result, err_msg = ProjectCodePrdService._request_markdown(prompt, max_tokens=2500)
        if err_msg:
            return '', err_msg
        text = ProjectCodePrdService._strip_markdown_fence(result)
        if not text.strip():
            return '', f'第{chunk_index}个代码分析agent未返回内容'
        return text.strip(), ''

    @staticmethod
    def _split_summary_chunks(summary, max_chars=9000):
        summary = summary or ''
        parts = re.split(r'(?=\n### )', summary)
        header = parts[0] if parts else ''
        file_parts = parts[1:] if len(parts) > 1 else []
        if not file_parts and len(summary) <= max_chars:
            return [summary]
        if not file_parts:
            return [summary[i:i + max_chars] for i in range(0, len(summary), max_chars)]
        chunks = []
        current = header.strip()
        for part in file_parts:
            candidate = (current + '\n' + part).strip() if current else part.strip()
            if current and len(candidate) > max_chars:
                chunks.append(current.strip())
                current = (header + '\n' + part).strip() if len(part) < max_chars else part.strip()
                continue
            current = candidate
        if current.strip():
            chunks.append(current.strip())
        return chunks or [summary]

    @staticmethod
    def build_docx(markdown):
        try:
            from docx import Document
        except Exception:
            return None, '未安装 python-docx，暂无法导出 docx'
        doc = Document()
        lines = (markdown or '').splitlines()
        in_code = False
        code_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('```'):
                if in_code:
                    if code_lines:
                        doc.add_paragraph('\n'.join(code_lines))
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                continue
            if in_code:
                code_lines.append(line)
                continue
            if stripped.startswith('#'):
                level = min(len(stripped) - len(stripped.lstrip('#')), 4)
                title = stripped.lstrip('#').strip()
                if level == 1:
                    doc.add_heading(title, level=1)
                else:
                    doc.add_heading(title, level=level)
            elif stripped.startswith(('- ', '* ')):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+[\.、]\s+', stripped):
                doc.add_paragraph(re.sub(r'^\d+[\.、]\s+', '', stripped), style='List Number')
            elif stripped:
                doc.add_paragraph(stripped)
            else:
                doc.add_paragraph('')
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio, ''

    @staticmethod
    def extract_title(markdown):
        for line in (markdown or '').splitlines():
            if line.startswith('# '):
                return line[2:].strip()[:256]
        return ''

    @staticmethod
    def _request_markdown(prompt, max_tokens=12000):
        try:
            from openai import OpenAI
            from config.ai_config import AIConfig
            import httpx
            import time
            api_key = AIConfig.get_api_key()
            api_base = AIConfig.get_api_base()
            model = AIConfig.get_model()
            key_source = AIConfig.get_api_key_source()
            if not api_key or api_key == '请替换为你的Meteor API Key':
                return '', '未配置API密钥，请在.env中配置METEOR_API_KEY'
            is_plan_key = '/plan/' in api_base
            request_base = AIService._normalize_plan_api_base(api_base) if is_plan_key else AIService._normalize_api_base(api_base)
            read_timeout = max(AIConfig.READ_TIMEOUT, 300)
            timeout = httpx.Timeout(connect=AIConfig.CONNECT_TIMEOUT, read=read_timeout, write=read_timeout, pool=AIConfig.CONNECT_TIMEOUT)
            logger.info(f'代码PRD AI配置: provider={AIConfig.MODEL_PROVIDER}, base={request_base}, model={model}, key_source={key_source}, key_prefix={api_key[:8]}, plan_key={is_plan_key}')
            gateway_retry_count = 0
            gateway_retry_delay = max(1.0, getattr(AIConfig, 'GATEWAY_RETRY_DELAY', 30.0))
            gateway_max_retries = max(0, getattr(AIConfig, 'GATEWAY_MAX_RETRIES', 0))
            while True:
                try:
                    result = AIService._request_model(
                        OpenAI, AIConfig, api_key, request_base, model, is_plan_key, prompt, timeout, httpx,
                        max_retries=1,
                        max_tokens=max_tokens,
                        temperature=0.2,
                        system_prompt='你是资深产品经理和软件架构分析师。根据代码库摘要输出严谨的PRD Markdown。'
                    )
                    return result, ''
                except Exception as request_err:
                    err_str = str(request_err)
                    gateway_error = AIService._is_gateway_error(err_str)

                    retry_available = gateway_max_retries == 0 or gateway_retry_count < gateway_max_retries
                    if gateway_error and retry_available:
                        gateway_retry_count += 1
                        logger.warning(f'代码PRD AI网关超时，等待{gateway_retry_delay}秒后继续重试，retry={gateway_retry_count}: {err_str[:200]}')
                        time.sleep(gateway_retry_delay)
                        continue
                    raise
        except Exception as e:
            logger.error(f'AI生成代码PRD失败: {e}')
            return '', f'AI生成代码PRD失败: {e}'


    @staticmethod
    def _build_prd_chunk_prompt(chunk, chunk_index, total_chunks, repo_url, branch, prompt_append=''):
        prompt_append = (prompt_append or '').strip()
        prompt_append_block = f'''\n<user-prompt-append>\n{prompt_append}\n</user-prompt-append>\n''' if prompt_append else ''
        return f'''
你是代码库 PRD 分析 agent。请只分析当前代码摘要分块，输出结构化 Markdown，不要生成最终 PRD。

<repo>
仓库地址：{repo_url}
分支：{branch}
分块：{chunk_index}/{total_chunks}
</repo>

<code-summary-chunk>
{chunk}
</code-summary-chunk>
{prompt_append_block}
输出要求：
1. 只输出 Markdown。
2. 按模块归纳：业务能力、用户角色、关键页面/接口/任务、核心流程、数据对象、权限点、异常边界、待确认项。
3. 只写当前分块有代码依据的结论，不要编造。
4. 每个结论尽量标注相关文件路径。
5. 控制篇幅，避免逐行解释代码。
'''.strip()

    @staticmethod
    def _build_prd_prompt(summary, repo_url, branch, prompt_append=''):
        prompt_append = (prompt_append or '').strip()
        prompt_append_block = f'''\n<user-prompt-append>\n{prompt_append}\n</user-prompt-append>\n''' if prompt_append else ''
        return f'''
请使用 gitnexus-exploring 和 product-design 的分析思路，根据代码库分块agent分析结果深度分析并输出最终 PRD 文档。

<repo>
仓库地址：{repo_url}
分支：{branch}
</repo>

<agent-analysis-summary>
{summary}
</agent-analysis-summary>
{prompt_append_block}
输出要求：
1. 只输出 Markdown，不要解释过程。
2. PRD 必须包含：背景与目标、用户角色、范围、功能清单、核心业务流程、详细需求、数据对象、权限、异常与边界、埋点/日志、验收标准、风险与待确认项。
3. “核心业务流程”中每个流程都必须补充完整 Mermaid 图，按流程性质选择 flowchart 或 sequenceDiagram。
4. Mermaid 图必须放在 ```mermaid 代码块中。
5. 对每个流程给出：触发条件、参与角色、前置条件、主流程、异常分支、后置结果。
6. 结合 product-design 视角补充页面/交互/信息架构建议，但不要写营销页内容。
7. 不要编造分块agent分析中没有依据的外部系统；不确定处放入“待确认项”。
'''.strip()

    @staticmethod
    def _scan_repo(repo_dir, repo_url, branch):
        file_infos = []
        tree_lines = []
        max_files = 80
        max_file_chars = 600
        ignore_dirs = {'.git', 'node_modules', 'dist', 'build', '__pycache__', '.venv', 'venv', 'target', '.idea', '.vscode'}
        include_exts = {'.py', '.js', '.ts', '.tsx', '.vue', '.java', '.go', '.rs', '.php', '.md', '.yml', '.yaml', '.json', '.sql'}
        for root, dirs, files in os.walk(repo_dir):
            dirs[:] = [d for d in dirs if d not in ignore_dirs]
            rel_root = os.path.relpath(root, repo_dir)
            depth = 0 if rel_root == '.' else rel_root.count(os.sep) + 1
            if depth <= 2:
                tree_lines.append(rel_root if rel_root != '.' else '.')
            for filename in files:
                ext = os.path.splitext(filename)[1].lower()
                if ext not in include_exts:
                    continue
                path = os.path.join(root, filename)
                rel_path = os.path.relpath(path, repo_dir).replace('\\', '/')
                if len(file_infos) >= max_files:
                    continue
                try:
                    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read(max_file_chars)
                except Exception:
                    text = ''
                file_infos.append({
                    'path': rel_path,
                    'head': ProjectCodePrdService._compact_text(text)
                })
        files_text = '\n\n'.join([f"### {item['path']}\n{item['head']}" for item in file_infos])
        return f'''
仓库：{repo_url}
分支：{branch}
目录概览：
{chr(10).join(tree_lines[:120])}

关键文件摘录（最多{max_files}个）：
{files_text}
'''.strip()

    @staticmethod
    def _safe_path_name(name):
        safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '-', str(name or '').strip())
        safe = re.sub(r'\s+', ' ', safe).strip(' .')
        return safe[:80] or 'default'

    @staticmethod
    def _compact_text(text):
        lines = []
        for line in (text or '').splitlines():
            s = line.rstrip()
            if not s:
                continue
            lines.append(s[:240])
            if len(lines) >= 40:
                break
        return '\n'.join(lines)

    @staticmethod
    def _strip_markdown_fence(text):
        s = (text or '').strip()
        m = re.match(r'^```(?:markdown|md)?\s*([\s\S]*?)\s*```$', s)
        return m.group(1).strip() if m else s

    @staticmethod
    def _is_safe_git_url(repo_url):
        s = (repo_url or '').strip()
        if not s or any(ch in s for ch in ['\n', '\r', ';', '&', '|', '`', '$', '<', '>']):
            return False
        return bool(re.match(r'^(https?://|git@)[\w.\-/:~]+(?:\.git)?$', s))
